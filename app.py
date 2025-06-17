from flask import Flask, render_template, send_file
from crawler import fetch_news
from datetime import datetime
from docx import Document
import os

app = Flask(__name__)
news_cache = {}

@app.route('/')
def index():
    global news_cache
    news_cache = fetch_news()
    return render_template('index.html', news=news_cache)

@app.route('/export')
def export_news():
    doc = Document()
    doc.add_heading('Threat News Report (India)', level=1)
    for date, items in news_cache.items():
        doc.add_heading(date, level=2)
        for item in items:
            doc.add_paragraph(f"{item['title']}", style='List Bullet')
            doc.add_paragraph(f"Source: {item['source']}", style='Intense Quote')
            doc.add_paragraph(f"Medium: {item['medium']}", style='Intense Quote')
            doc.add_paragraph(f"Link: {item['link']}")
            doc.add_paragraph()

    filename = f"news_report_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    filepath = os.path.join("exports", filename)
    os.makedirs("exports", exist_ok=True)
    doc.save(filepath)
    return send_file(filepath, as_attachment=True)

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(host='0.0.0.0', port=port)

